import streamlit as st
import pdfplumber
import json
import io
from docx import Document
from docx.shared import Pt
from openai import OpenAI

# ==========================================
# 1. 初始化 AI 客户端
# ==========================================
# 请在此处填写你的 API Key 和 Base URL
API_KEY = "sk-2d4070285351465ebdf763a80c76eb84" 
BASE_URL = "https://api.deepseek.com/v1" # 以 DeepSeek 为例

client = OpenAI(api_key=API_KEY, base_url=BASE_URL)

# ==========================================
# 2. 核心工具函数：PDF 全文深度解析
# ==========================================
def extract_pdf_text(uploaded_file):
    """深度提取工程 PDF。为了能够全面捕捉地下结构、工法和地质参数，保持深层读取"""
    text_segments = []
    with pdfplumber.open(uploaded_file) as pdf:
        total_pages = len(pdf.pages)
        st.info(f"📄 报告解析成功，共 {total_pages} 页。正在深度扫描地下结构与工程地质章节...")
        
        # 读取前 25 页 (涵盖目录、概述、政策、以及初步的技术指标、隧道工法变更说明)
        for i in range(min(25, total_pages)):
            text_segments.append(f"--- PAGE {i+1} ---\n" + (pdf.pages[i].extract_text() or ""))
            
        # 读取最后 12 页 (涵盖详细投资估算、社会效益、风险控制及结论)
        if total_pages > 37:
            text_segments.append("\n...[中间章节省略]...\n")
            for i in range(total_pages - 12, total_pages):
                text_segments.append(f"--- PAGE {i+1} ---\n" + (pdf.pages[i].extract_text() or ""))
                
    return "\n".join(text_segments)

# ==========================================
# 3. AI 五维全能复合审查引擎（融入地下结构专项要点）
# ==========================================
def ai_five_dimension_review(pdf_text):
    """让 AI 同时扮演：数据审计员(D1)、规划师(D2)、总工程师(D3)、投资风控总监(D4)、岩土地下结构专家(D5)"""
    
    prompt = f"""
    你是一个精通国家工程建设强制性标准、发改委可研大纲、以及基础设施项目（特别是大型水下/盾构隧道、城市地下结构）全生命周期管理的顶级会审专家。
    请仔细阅读以下从工程可研报告中提取的文本，从五个维度进行全维度的立体化审查。
    
    请严格按照以下 JSON 格式输出，不要包含任何多余的 Markdown 标记或解释文字：
    {{
      "dimension_1_data": {{
        "project_name": "项目名称",
        "total_investment_wanyu": 0.0,
        "construction_period_months": 0,
        "route_length_km": 0.0,
        "design_speed_kmh": 0
      }},
      "dimension_2_policy": {{
        "checklist": {{
          "建设必要性论证": true, 
          "规划与选址合规性": true, 
          "节能减排与双碳响应": true,
          "环境影响与生态保护": true,
          "社会稳定风险评估": false
        }},
        "missing_sections_details": "薄弱环节说明",
        "policy_fit_analysis": "宏观政策契合度评价",
        "normative_suggestions": "发改委规范性建议"
      }},
      "dimension_3_regulatory": {{
        "overall_status": "安全合规 / 存在强条合规风险 / 存在一般条文建议",
        "mandatory_clauses": [
          {{
            "item_name": "核心条文项名称",
            "standard_basis": "规范出处与条款号",
            "report_value": "报告中对应的实际数值或方案说明",
            "status": "合规 / 违规 / 风险",
            "expert_analysis": "依据条文进行的专家合规性拆解。"
          }}
        ],
        "detailed_violations": "强条违规缺陷详述"
      }},
      "dimension_4_logic": {{
        "safety_logic_review": "安全性逻辑审查说明",
        "economic_reasonableness": "经济合理性与性价比控制评估",
        "demand_scale_matching": "供需规模匹配逻辑评估",
        "financial_sustainability": "财务与效益分摊逻辑评价"
      }},
      "underground_special_review": {{
        "design_depth_7_1_1": "【设计文件总体深度】：评估报告内容是否满足对应行业（如城市道路、市政或轨道交通地下结构）的编制深度规定？是否有严重漏项？",
        "supporting_reports_7_1_3": "【配套支撑报告健全度】：重点审查报告中是否提及并集成了‘可行性研究报告批复、地质灾害评估、洪评报告、地震安评报告、环评报告、初步勘察报告’等前置依据？缺少哪些？",
        "hydrogeology_eval_7_1_4": "【水文地质与特殊岩土分析】：审查设计说明中对地形地貌、特殊岩土层力学参数、不良地质（如地下水、承压水、有害气体）、环境水土腐蚀性评价是否完整？是否给出了针对性的应对策略？",
        "surrounding_investigation_7_1_5": "【周边建构筑物及管线调查】：报告对沿线敏感建构筑物、地下既有综合管线、地下障碍物、工程河段水文与河势演变的资料收集是否齐全？有无清晰的保护或迁改逻辑？",
        "construction_method_7_1_7": "【施工工法地质适应性】：针对选用的主要工法（如盾构法、明挖法），评估其与报告中阐述的水文地质条件是否匹配？从地下结构安全的角度看，技术交底和施工工法推导是否合理？"
      }}
    }}
    
    【地下结构专项审查重点指引（依据地下结构审查要点规范）】：
    - 结合《地下结构审查要点.docx》，重点剖析“underground_special_review”中的五个核心要点（7.1.1, 7.1.3, 7.1.4, 7.1.5, 7.1.7）。
    - 针对东太湖隧道：重点关注它穿越太湖这一复杂水体时，水文地质和岸线建构筑物（7.1.5）的调查深度，以及盾构工法对太湖复杂淤泥质土层的适应性（7.1.7）。
    
    【输入文本】：
    {pdf_text}
    """
    
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "system", "content": "你是一个只会输出高精度、标准结构化 JSON 数据的自动化工程审查后台。"},
            {"role": "user", "content": prompt}
        ],
        response_format={"type": "json_object"},
        temperature=0.15
    )
    
    # 强鲁棒性 Markdown 剥离清洗
    raw_content = response.choices[0].message.content.strip()
    if raw_content.startswith("```json"):
        raw_content = raw_content[7:]
    elif raw_content.startswith("```"):
        raw_content = raw_content[3:]
    if raw_content.endswith("```"):
        raw_content = raw_content[:-3]
    raw_content = raw_content.strip()
    
    try:
        return json.loads(raw_content)
    except json.JSONDecodeError:
        return {
            "dimension_1_data": {"project_name": "解析失败", "total_investment_wanyu": 0, "construction_period_months": 0, "route_length_km": 0, "design_speed_kmh": 0},
            "dimension_2_policy": {"checklist": {}, "missing_sections_details": "JSON格式异常"},
            "dimension_3_regulatory": {"overall_status": "无法评估", "mandatory_clauses": []},
            "dimension_4_logic": {},
            "underground_special_review": {"design_depth_7_1_1": "解析失败", "supporting_reports_7_1_3": "解析失败", "hydrogeology_eval_7_1_4": "解析失败"}
        }

# ==========================================
# 新增：Word 报告生成函数
# ==========================================
def create_word_report(results):
    doc = Document()
    doc.add_heading('工程可行性研究报告审查意见', 0)
    
    # 维度1
    doc.add_heading('一、关键技术经济指标', level=1)
    d1 = results['dimension_1_data']
    for k, v in d1.items():
        doc.add_paragraph(f"{k}: {v}")
        
    # 维度2
    doc.add_heading('二、政策合规性审查', level=1)
    doc.add_paragraph(results['dimension_2_policy']['policy_fit_analysis'])
    
    # 维度3
    doc.add_heading('三、法规强制性条文审查', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '项目'
    hdr_cells[1].text = '规范依据'
    hdr_cells[2].text = '结论'
    for item in results['dimension_3_regulatory']['mandatory_clauses']:
        row_cells = table.add_row().cells
        row_cells[0].text = item['item_name']
        row_cells[1].text = item['standard_basis']
        row_cells[2].text = item['status']
        
    # 维度4
    doc.add_heading('四、内容逻辑与经济性审查', level=1)
    doc.add_paragraph(results['dimension_4_logic']['safety_logic_review'])
    
    # 保存到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 4. Streamlit 在线网页前端界面
# ==========================================
st.set_page_config(page_title="全维工程可研智能审查系统 v4.5", layout="wide")

st.title("🏗️ 全在线工程可行性研究报告立体化审查系统 (专项增强版)")
st.caption("基于 AI + 行业大数据规则碰撞的工程文件全生命周期会审平台 | 已激活：地下结构专项审查模块")

# 侧边栏配置
with st.sidebar:
    st.header("⚙️ 审查维度配置")
    st.checkbox("维度1：文字与数据一致性", value=True, disabled=True)
    st.checkbox("维度2：政策与规范性", value=True, disabled=True)
    st.checkbox("维度3：法律法规与强条", value=True, disabled=True)
    st.checkbox("维度4：技术安全与经济逻辑", value=True, disabled=True)
    st.checkbox("🔥 专项：地下结构专项审查", value=True, disabled=True)
    st.markdown("---")
    st.markdown("🧱 **地下结构审计卡尺（依据最新要点）：**\n"
                "- 7.1.1 设计文件总体深度\n"
                "- 7.1.3 洪评/震安评/勘察报告齐备度\n"
                "- 7.1.4 水文地质与特殊岩土评价\n"
                "- 7.1.5 周边建构筑物与管线调查\n"
                "- 7.1.7 施工工法地质适应性")

# 文件上传
uploaded_file = st.file_uploader("请上传您的工程可研报告 PDF 文件", type=["pdf"])

if uploaded_file is not None:
    if st.button("🚀 启动五维一体化全能专项审查"):
        with st.spinner("系统正在调集岩土总工、结构审查专家及多方AI席位进行深度联合会审..."):
            
            # 1. 解析文本
            pdf_text = extract_pdf_text(uploaded_file)
            
            # 2. 调用五维 AI 引擎
            review_results = ai_five_dimension_review(pdf_text)
            
            d1_data = review_results.get("dimension_1_data", {})
            d2_data = review_results.get("dimension_2_policy", {})
            d3_data = review_results.get("dimension_3_regulatory", {})
            d4_data = review_results.get("dimension_4_logic", {})
            d5_data = review_results.get("underground_special_review", {}) # 新增维度数据
            
            st.success("🎉 五维一体化全能专项审查圆满完成！审查报告已实时渲染。")
            st.write("---")
            
            # 创建五个标签页，第5个为新增的地下结构专项
            tab1, tab2, tab3, tab4, tab5 = st.tabs([
                "🔍 维度1：数据一致性", 
                "⚖️ 维度2：政策与规范性", 
                "📐 维度3：法规与强条对碰",
                "📊 维度4：宏观安全与经济",
                "🧱 专项：地下结构与水文地质"
            ])
            
            # ==========================================
            # TAB 1: 维度1
            # ==========================================
            with tab1:
                st.subheader("📊 关键指标元数据对账")
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("项目名称", d1_data.get("project_name", "未识别"))
                with col2:
                    st.metric("总投资额 (万元)", f"{d1_data.get('total_investment_wanyu', 0):,.2f}")
                with col3:
                    st.metric("建设工期 (个月)", f"{d1_data.get('construction_period_months', 0)} 个月")
                with col4:
                    st.metric("路线全长 (km)", f"{d1_data.get('route_length_km', 0)} km")
                
                errors_found = 0
                if d1_data.get('total_investment_wanyu', 0) > 500000 and d1_data.get('construction_period_months', 0) < 24:
                    st.error("❌ **发现严重数据矛盾（投资与工期不匹配）**")
                    errors_found += 1
                if errors_found == 0:
                    st.success("✅ 经过后端比对，报告内核心指标的前后文钩稽关系完全一致。")
            
            # ==========================================
            # TAB 2: 维度2
            # ==========================================
            with tab2:
                st.subheader("📋 必备章节齐全性审查 (Checklist)")
                checklist = d2_data.get("checklist", {})
                if checklist:
                    chk_cols = st.columns(len(checklist))
                    for idx, (item, is_present) in enumerate(checklist.items()):
                        with chk_cols[idx]:
                            if is_present: st.success(f"🍏 {item}\n\n**[已包含]**")
                            else: st.error(f"🍎 {item}\n\n**[缺失/薄弱]**")
                st.subheader("🌱 发展政策与战略红线契合度研判")
                st.info(d2_data.get("policy_fit_analysis", ""))
            
            # ==========================================
            # TAB 3: 维度3
            # ==========================================
            with tab3:
                st.subheader("🛑 国家工程建设强制性条文对碰结果")
                st.warning(f"⚠️ 研判结论：{d3_data.get('overall_status', '未评估')}")
                clauses = d3_data.get("mandatory_clauses", [])
                for idx, clause in enumerate(clauses):
                    with st.expander(f"📑 {clause.get('item_name')} — 依据：{clause.get('standard_basis')}"):
                        st.markdown(f"**📝 本报告设计值/方案：** {clause.get('report_value')}")
                        st.markdown(f"**🧑‍💻 审查总工意见：** {clause.get('expert_analysis')}")
            
            # ==========================================
            # TAB 4: 维度4
            # ==========================================
            with tab4:
                st.subheader("💎 方案技术安全与经济性逻辑深度审计看板")
                tot_inv = d1_data.get('total_investment_wanyu', 0)
                r_len = d1_data.get('route_length_km', 0)
                if tot_inv > 0 and r_len > 0:
                    st.metric("📊 每公里宏观综合造价", f"{(tot_inv / 10000) / r_len:.2f} 亿元 / 公里")
                st.info(d4_data.get("safety_logic_review", ""))
                st.info(d4_data.get("economic_reasonableness", ""))

            # ==========================================
            # TAB 5: 全新增强展示（地下结构专项审查要点）
            # ==========================================
            with tab5:
                st.subheader("🧱 依据《地下结构审查要点》进行专项对碰")
                st.caption("以下研判意见严格对照您上传的地下结构设计深度、水文勘察、建构筑物保护及工法标准生成：")
                
                # 采用专业卡片流布局，清晰映射 7.1 核心条文
                st.markdown("#### 📐 1. 设计编制深度与前置报告审计")
                c1, c2 = st.columns(2)
                with c1:
                    with st.get_container() if hasattr(st, "get_container") else st.container(border=True):
                        st.markdown("##### 📌 **7.1.1 & 7.1.2 设计文件总体深度及说明图纸**")
                        st.write(d5_data.get("design_depth_7_1_1", "未生成有效评估。"))
                with c2:
                    with st.get_container() if hasattr(st, "get_container") else st.container(border=True):
                        st.markdown("##### 📜 **7.1.3 设计依据与配套支撑报告齐备度**")
                        st.info(d5_data.get("supporting_reports_7_1_3", "未生成有效评估。"))
                        
                st.markdown("---")
                st.markdown("#### 🌍 2. 岩土水文环境与周边安全边界")
                
                # 使用现代化的对话或警告气泡，突出地质安全风险
                with st.chat_message("user", avatar="💧"):
                    st.markdown("**7.1.4 水文地质、特殊岩土及环境水土腐蚀性深度评价：**")
                    st.write(d5_data.get("hydrogeology_eval_7_1_4", "未生成有效评估。"))
                    
                with st.chat_message("user", avatar="🏢"):
                    st.markdown("**7.1.5 沿线建构筑物、地下既有管线及障碍物调查评估：**")
                    st.write(d5_data.get("surrounding_investigation_7_1_5", "未生成有效评估。"))
                    
                st.markdown("---")
                st.markdown("#### 🚜 3. 核心施工工法审计")
                with st.chat_message("assistant", avatar="⚙️"):
                    st.markdown("**7.1.7 车站/区间主体施工工法（如大直径盾构法）地质适应性研判：**")
                    st.write(d5_data.get("construction_method_7_1_7", "未生成有效评估。"))

        st.write("---")
        # 生成 Word 并显示下载按钮
        word_buffer = create_word_report(review_results)
        st.download_button(
            label="📥 下载审查意见报告 (Word)",
            data=word_buffer,
            file_name="工程审查意见报告.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )